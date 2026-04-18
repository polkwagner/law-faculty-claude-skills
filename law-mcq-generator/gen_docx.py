"""Generate MCQ .docx deliverables from markdown draft sources.

Generates three documents from the MCQ skill's markdown drafts:
1. Exam questions (.docx) from draft_full_set.md
2. Full answer key (.docx) from draft_answer_key_full.md
3. Student answer key (.docx) from draft_answer_key_student.md

Uses the MCQ skill's templates (exam_template.docx, answer_key_template.docx,
student_answer_key_template.docx) for consistent formatting.

Usage:
    python gen_docx.py <source_dir> [--output <output_dir>] [--prefix <filename_prefix>]
                       [--school <school>] [--course <course>] [--semester <semester>]
                       [--professor <professor>]

    source_dir: directory containing draft_full_set.md, draft_answer_key_full.md,
                and draft_answer_key_student.md

    Defaults:
      --output     ~/Downloads/
      --prefix     IP_Final_Exam_MCQ
      --school     UNIVERSITY OF PENNSYLVANIA CAREY LAW SCHOOL
      --course     INTRODUCTION TO INTELLECTUAL PROPERTY — LAW 507
      --semester   SPRING 2026
      --professor  [Your Name]
"""

import argparse
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

# ── Paths ──

HOME = os.path.expanduser("~")
SKILL_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)))
DOWNLOADS = os.path.join(HOME, "Downloads")


def clear_placeholders(doc):
    """Remove all placeholder paragraphs from template."""
    for p in doc.paragraphs[:]:
        doc.element.body.remove(p._element)


def apply_numbering(paragraph, num_id, ilvl=0):
    """Attach <w:numPr numId=N ilvl=L/> to a paragraph, overriding style defaults."""
    pPr = paragraph._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


# abstractNumId values that exist in exam_template.docx/numbering.xml
ABSTRACT_NUM_DECIMAL = 7   # "1.", "2.", "3." ordered list — used for instructions
ABSTRACT_NUM_LOWER_ALPHA = 8  # "(a)", "(b)", "(c)", "(d)" — used for Answer style

# numId the template uses for the instruction ordered list (1., 2., 3., ...)
INSTRUCTION_NUM_ID = 82


def register_answer_numids(doc, count):
    """Create `count` fresh numId entries in numbering.xml that reference the
    Answer-letter abstract (abstractNumId=8) with startOverride=1.

    The MCQ exam has N questions each with four (a)(b)(c)(d) choices. Word's
    list numbering is stateful: if every question reused the same numId, the
    answer letters would run continuously — so Q2 would render as (e)(f)(g)(h).
    The fix: give each question its OWN numId that starts fresh at (a). The
    template ships ~20 reset-numIds (85, 87–105), not enough for 40 questions,
    so we synthesize new ones at runtime.

    Returns the list of numIds created.
    """
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId")))
                for n in numbering.findall(qn("w:num"))]
    next_id = (max(existing) if existing else 0) + 1

    new_ids = []
    for i in range(count):
        nid = next_id + i
        num_xml = (
            '<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:numId="{nid}">'
            f'<w:abstractNumId w:val="{ABSTRACT_NUM_LOWER_ALPHA}"/>'
            '<w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'
            '</w:num>'
        )
        numbering.append(parse_xml(num_xml))
        new_ids.append(nid)
    return new_ids


def add_hr(doc):
    """Add horizontal rule (gray bottom border paragraph)."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "808080")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    pPr.append(spacing)


# ═══════════════════════════════════════════════════════════════
# Document 1: Exam
# ═══════════════════════════════════════════════════════════════

def generate_exam(exam_md, output_path, school, course, semester, professor,
                   header_text=None):
    """Generate exam .docx from draft_full_set.md.

    Args:
        exam_md: path to draft_full_set.md
        output_path: path to write the .docx
        school: school name for title page
        course: course name for title page
        semester: semester/year for title page
        professor: professor name for title page
        header_text: optional header override (default: derived from course/semester)
    """
    doc = Document(os.path.join(SKILL_DIR, "exam_template.docx"))
    clear_placeholders(doc)

    # Pre-allocate one fresh (a)(b)(c)(d) numbering instance per question so
    # each question's answer letters restart at (a). 60 is a safe ceiling.
    fresh_answer_numids = register_answer_numids(doc, 60)
    question_index = 0  # incremented each time we open a new question

    # Update header. Default layout: three-part (left/center/right) using a
    # center-aligned tab stop at the midpoint of the content area and a
    # right-aligned tab stop at the right margin — e.g.
    #   [ LAW 507        INTRODUCTION TO INTELLECTUAL PROPERTY        SPRING 2026 ]
    # If the caller supplies a literal --header string, honor it as-is.
    use_three_part = not header_text
    if use_three_part:
        m = re.search(r'(.+?)\s*[—–-]\s*(LAW\s+\S+)', course, re.IGNORECASE)
        if m:
            course_name = m.group(1).strip()
            course_code = m.group(2).strip()
        else:
            course_name, course_code = course.strip(), ""
        header_text = f"[ {course_code}\t{course_name}\t{semester} ]"

    for section in doc.sections:
        content_width = section.page_width - section.left_margin - section.right_margin
        for hp in section.header.paragraphs:
            if not hp.text.strip():
                continue
            for run in hp.runs:
                run.text = ""
            hp.runs[0].text = header_text
            if use_three_part:
                hp.paragraph_format.tab_stops.clear_all()
                hp.paragraph_format.tab_stops.add_tab_stop(
                    content_width // 2, WD_TAB_ALIGNMENT.CENTER)
                hp.paragraph_format.tab_stops.add_tab_stop(
                    content_width, WD_TAB_ALIGNMENT.RIGHT)
            break

    with open(exam_md) as f:
        md = f.read()

    # Parse the markdown
    lines = md.split("\n")
    i = 0

    # ── Title page ──
    p = doc.add_paragraph(style="First Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(school)
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(course)
    run.italic = True

    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FINAL EXAMINATION — {semester}")

    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(professor)

    doc.add_paragraph(style="Body Text")

    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MULTIPLE CHOICE QUESTIONS")
    run.bold = True

    doc.add_paragraph(style="Body Text")

    # ── Parse and add content ──
    # Skip any leading blank lines or title lines in the md
    while i < len(lines) and not lines[i].strip():
        i += 1

    # Skip the markdown's own title block — the title page is built
    # programmatically above from school/course/semester/professor, so any
    # leading title lines in the draft (school name, course name, semester,
    # professor, "PART 1: ...") would otherwise be re-emitted as body text.
    # Fast-forward to the first real instruction paragraph.
    title_patterns = [
        re.compile(r'^\*{0,2}UNIVERSITY OF .+\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}INTRODUCTION TO .+\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}FINAL EXAM.*\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}PRACTICE .*QUESTIONS\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}SPRING \d{4}\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}FALL \d{4}\*{0,2}$', re.IGNORECASE),
        re.compile(r'^\*{0,2}Professor .+\*{0,2}$'),
        re.compile(r'^\*{0,2}PART \d+:.*\*{0,2}$', re.IGNORECASE),
    ]
    def is_title_line(line):
        stripped = line.strip()
        return any(p.match(stripped) for p in title_patterns)

    # Consume consecutive title/blank lines at the top of the md
    while i < len(lines) and (not lines[i].strip() or is_title_line(lines[i])):
        i += 1

    # Find and add instruction paragraphs (before first fact pattern)
    instruction_lines = []
    while i < len(lines):
        line = lines[i]
        if re.match(r'\[ Questions \d+', line):
            break
        instruction_lines.append(line)
        i += 1

    # Process instruction text
    instr_text = "\n".join(instruction_lines).strip()
    if instr_text:
        # Split into paragraphs
        for para_text in instr_text.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            # Check if it's a bullet point
            if para_text.startswith("- ") or para_text.startswith("• "):
                for bullet_line in para_text.split("\n"):
                    bullet_line = bullet_line.lstrip("- •").strip()
                    if bullet_line:
                        p = doc.add_paragraph(style="List Paragraph")
                        # Handle bold within bullets
                        parts = re.split(r'(\*\*.*?\*\*)', bullet_line)
                        for part in parts:
                            if part.startswith("**") and part.endswith("**"):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
            else:
                # Regular paragraph — rendered as a numbered instruction
                # (1., 2., 3., ...) using the template's decimal list (numId 82).
                clean = para_text.replace("\n", " ")
                p = doc.add_paragraph(style="Body Text")
                apply_numbering(p, INSTRUCTION_NUM_ID, ilvl=0)
                parts = re.split(r'(\*\*.*?\*\*)', clean)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    # ── Fact patterns and questions ──
    while i < len(lines):
        line = lines[i].strip()

        # Fact pattern header: [ Questions X through Y relate to Fact Pattern Z ]
        fp_match = re.match(r'\[ Questions (\d+) through (\d+) relate to Fact Pattern ([A-G]) \]', line)
        if fp_match:
            q_start, q_end, fp_letter = fp_match.group(1), fp_match.group(2), fp_match.group(3)

            # Each fact pattern starts on a new page. page_break_before on
            # the paragraph is preferable to an inline page-break run because
            # it moves with the paragraph if content before it shifts.
            header_text = line.strip("[] ").strip()
            p = doc.add_paragraph(style="Body Text")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = True
            p.add_run(header_text)

            # Next: FACT PATTERN X title
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1

            # Title: FACT PATTERN X (might be in markdown as bold or plain)
            if i < len(lines):
                title_line = lines[i].strip().replace("**", "")
                if title_line.startswith("FACT PATTERN"):
                    p = doc.add_paragraph(style="Title")
                    p.add_run(title_line)
                    i += 1

            # Subtitle: The one with the...
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                sub_line = lines[i].strip().replace("*", "")
                if sub_line.lower().startswith("the one with"):
                    p = doc.add_paragraph(style="Subtitle")
                    p.add_run(sub_line)
                    i += 1

            # Empty line after subtitle
            doc.add_paragraph(style="Body Text")

            # Narrative paragraphs (until first question)
            narrative_lines = []
            while i < len(lines):
                line = lines[i].strip()
                # Check if we hit a question (starts with **N. )
                if re.match(r'\*\*\d+\.', line):
                    break
                # Check if we hit next fact pattern
                if re.match(r'\[ Questions \d+', line):
                    break
                # Check for end marker
                if line.startswith("[ END"):
                    break
                narrative_lines.append(lines[i])
                i += 1

            # Add narrative paragraphs
            nar_text = "\n".join(narrative_lines).strip()
            for para in nar_text.split("\n\n"):
                para = para.strip()
                if para:
                    p = doc.add_paragraph(style="Body Text")
                    # Handle italic (case names)
                    parts = re.split(r'(\*[^*]+\*)', para.replace("\n", " "))
                    for part in parts:
                        if part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        else:
                            p.add_run(part)

            # Questions
            while i < len(lines):
                line = lines[i].strip()
                # New fact pattern or end
                if re.match(r'\[ Questions \d+', line) or line.startswith("[ END"):
                    break

                # Question stem: **N. text**
                q_match = re.match(r'\*\*(\d+)\.\s*(.*)', line)
                if q_match:
                    q_num = q_match.group(1)
                    # Collect full stem (may span multiple lines until answer choices)
                    stem_parts = [q_match.group(2)]
                    i += 1
                    while i < len(lines):
                        next_line = lines[i].strip()
                        # Stop at answer choices: > (a), (a), > (b), (b), etc.
                        if re.match(r'>?\s*\([a-d]\)', next_line):
                            break
                        if re.match(r'\*\*\d+\.', next_line):
                            break
                        if re.match(r'\[ Questions \d+', next_line) or next_line.startswith("[ END"):
                            break
                        if next_line:
                            stem_parts.append(next_line)
                        i += 1

                    stem_text = " ".join(stem_parts).rstrip("*").strip()

                    # Add question paragraph. The Question style handles
                    # indent/hanging layout; stem text is plain (not bolded)
                    # — only the italic spans (case names) stay italicized.
                    p = doc.add_paragraph(style="Question")
                    p.add_run(f"{q_num}. ")
                    p.add_run("\t")
                    stem_clean = stem_text.replace("**", "")
                    parts = re.split(r'(\*[^*]+\*)', stem_clean)
                    for part in parts:
                        if part.startswith("*") and part.endswith("*"):
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        else:
                            p.add_run(part)

                    question_index += 1
                    q_numid = fresh_answer_numids[question_index - 1]

                    # Answer choices
                    choices_found = 0
                    while i < len(lines) and choices_found < 4:
                        aline = lines[i].strip()
                        # Skip blank lines between choices
                        if not aline:
                            i += 1
                            continue
                        # Match > (a) text or (a) text (with optional tab after letter)
                        choice_match = re.match(r'>?\s*\(([a-d])\)\s*\t?(.*)', aline)
                        if choice_match:
                            letter = choice_match.group(1)
                            choice_text = choice_match.group(2).strip()

                            # Choice text may continue on next lines
                            i += 1
                            while i < len(lines):
                                next_line = lines[i].strip()
                                # Blank line = end of this choice
                                if not next_line:
                                    break
                                # Next choice
                                if re.match(r'>?\s*\([a-d]\)', next_line):
                                    break
                                if re.match(r'\*\*\d+\.', next_line):
                                    break
                                if re.match(r'\[ Questions \d+', next_line) or next_line.startswith("[ END"):
                                    break
                                # Continuation line
                                choice_text += " " + next_line.lstrip("> ").strip()
                                i += 1

                            choices_found += 1
                            p = doc.add_paragraph(style="Answer")
                            # All four choices share the same per-question
                            # numId. The first paragraph starts the list (at
                            # "a" via startOverride=1); subsequent paragraphs
                            # continue it (b, c, d). If we only tagged the
                            # first one, the rest would inherit the style's
                            # default numId (84) and its counter would bleed
                            # across questions.
                            apply_numbering(p, q_numid, ilvl=0)
                            parts = re.split(r'(\*[^*]+\*)', choice_text)
                            for part in parts:
                                if part.startswith("*") and part.endswith("*"):
                                    run = p.add_run(part[1:-1])
                                    run.italic = True
                                else:
                                    p.add_run(part)
                        else:
                            break
                else:
                    i += 1

        elif line.startswith("[ END"):
            p = doc.add_paragraph(style="Body Text")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)
            i += 1

            # Add remaining content (Report of Ambiguity section etc.)
            remaining = []
            while i < len(lines):
                remaining.append(lines[i])
                i += 1
            rem_text = "\n".join(remaining).strip()
            if rem_text:
                for para in rem_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        p = doc.add_paragraph(style="Body Text")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run(para.replace("\n", "\n"))
        else:
            i += 1

    doc.save(output_path)
    print(f"Exam saved to {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# Document 2: Full Answer Key
# ═══════════════════════════════════════════════════════════════

def generate_answer_key(ak_full_md, output_path, school, course, semester, professor):
    """Generate full answer key .docx from draft_answer_key_full.md."""
    doc = Document(os.path.join(SKILL_DIR, "answer_key_template.docx"))
    clear_placeholders(doc)

    with open(ak_full_md) as f:
        md = f.read()

    # ── Header block ──
    for text, bold in [
        (school, True),
        (f"{course}: MCQ ANSWER KEY", True),
        (semester, True),
        (professor, True),
        ("", False),
    ]:
        p = doc.add_paragraph(style="First Paragraph" if text == school else "Body Text")
        if text:
            run = p.add_run(text)
            run.bold = bold
        if text == school:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Glossary ──
    add_hr(doc)

    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KEY TO ANSWER KEY NOTATION")
    run.bold = True

    doc.add_paragraph(style="Body Text")

    glossary_sections = {
        "Cognitive Taxonomy Codes": [
            "EA = Element Application. Apply specific doctrinal elements or tests to facts.",
            "AE = Argument Evaluation. Identify which party has the stronger or best argument.",
            "FB = Factor Balancing. Weigh factors in a multi-factor test against ambiguous facts.",
            "RI = Regime Identification. Identify which legal framework, test, or body of law governs.",
            "DD = Doctrinal Distinction. Distinguish between related or easily confused doctrines.",
            "NR = Negative Recognition. Recognize when a doctrine does not apply despite surface similarity.",
        ],
        "Difficulty Estimates": [
            "M = Moderate. 70-85% of well-prepared students answer correctly.",
            "H = Hard. 40-65% of well-prepared students answer correctly.",
            "VH = Very Hard. 20-40% of well-prepared students answer correctly.",
        ],
        "Distractor Taxonomy Codes": [
            "CW = Correct Rule, Wrong Application. Right legal standard, misapplied to these facts.",
            "PA = Plausible Argument, Not the Law. Sounds right as policy but is not the doctrine.",
            "TN = True but Non-Responsive. Accurate legal statement that does not answer this question.",
            "IA = Incomplete Analysis. Gets part right but misses a critical element.",
            "CE = Common Student Error. Reflects a typical misconception or conflation.",
            "DC = Doctrine Confusion. Applies analysis from the wrong legal framework.",
            "SA = Superficially Attractive. Matches a surface feature but misses the deeper issue.",
        ],
    }

    for heading, entries in glossary_sections.items():
        p = doc.add_paragraph(style="Body Text")
        run = p.add_run(heading)
        run.bold = True
        for entry in entries:
            p = doc.add_paragraph(style="Body Text")
            run = p.add_run(entry)
            run.font.size = Pt(10)
        doc.add_paragraph(style="Body Text")

    add_hr(doc)

    # ── Parse questions from markdown ──
    # Split on "Question N" headers
    q_blocks = re.split(r'\n(?=Question \d+\n)', md)

    for block in q_blocks:
        block = block.strip()
        if not block.startswith("Question "):
            continue

        q_match = re.match(r'Question (\d+)\n', block)
        if not q_match:
            continue
        q_num = int(q_match.group(1))

        if q_num > 1:
            add_hr(doc)

        # Question N header
        p = doc.add_paragraph(style="First Paragraph")
        run = p.add_run(f"Question {q_num}")
        run.bold = True

        # Parse the rest line by line
        block_lines = block.split("\n")[1:]  # skip "Question N"

        in_distractor = False
        for bline in block_lines:
            bline_strip = bline.strip()
            if not bline_strip:
                continue

            # Distractor analysis entries: (a) [CODE]: text
            dist_match = re.match(r'\(([a-d])\)\s*\[(\w+)\]:\s*(.*)', bline_strip)
            if dist_match:
                letter, code, text = dist_match.groups()
                p = doc.add_paragraph(style="Compact")
                # Handle italic case names
                full_text = f"({letter}) [{code}]: {text}"
                parts = re.split(r'(\*[^*]+\*)', full_text)
                for part in parts:
                    if part.startswith("*") and part.endswith("*"):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
                continue

            # Explanation line (has bold lead)
            if bline_strip.startswith("Explanation:"):
                p = doc.add_paragraph(style="Body Text")
                run = p.add_run("Explanation: ")
                run.bold = True
                rest = bline_strip[len("Explanation:"):].strip()
                parts = re.split(r'(\*[^*]+\*)', rest)
                for part in parts:
                    if part.startswith("*") and part.endswith("*"):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
                continue

            # Distractor Analysis: header
            if bline_strip == "Distractor Analysis:":
                p = doc.add_paragraph(style="Body Text")
                run = p.add_run("Distractor Analysis:")
                run.bold = True
                continue

            # Regular metadata lines
            p = doc.add_paragraph(style="Body Text")
            parts = re.split(r'(\*[^*]+\*)', bline_strip)
            for part in parts:
                if part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    # ── Exam-level summary (everything after last question block) ──
    summary_match = re.search(r'\nEXAM-LEVEL QUALITY SUMMARY\n', md)
    if summary_match:
        add_hr(doc)
        summary_text = md[summary_match.start():].strip()
        for para in summary_text.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            # Check if it looks like a heading (all caps or short bold-like line)
            if para.isupper() or (len(para) < 50 and ":" in para):
                p = doc.add_paragraph(style="Body Text")
                run = p.add_run(para)
                run.bold = True
            else:
                for subline in para.split("\n"):
                    subline = subline.strip()
                    if subline:
                        p = doc.add_paragraph(style="Body Text")
                        p.add_run(subline)

    doc.save(output_path)
    print(f"Answer key saved to {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# Document 3: Student Answer Key
# ═══════════════════════════════════════════════════════════════

def generate_student_key(ak_student_md, output_path, school, course, semester, professor):
    """Generate student answer key .docx from draft_answer_key_student.md."""
    doc = Document(os.path.join(SKILL_DIR, "student_answer_key_template.docx"))
    clear_placeholders(doc)

    with open(ak_student_md) as f:
        md = f.read()

    # ── Header ──
    for text in [
        school,
        course,
        semester,
        professor,
        "ANSWER KEY FOR MULTIPLE CHOICE QUESTIONS",
    ]:
        p = doc.add_paragraph(style="Body Text")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True

    doc.add_paragraph(style="Normal")

    # ── Parse questions ──
    # Pattern: **Question N --- Correct Answer: (x)**
    q_blocks = re.split(r'\n(?=\*\*Question \d+)', md)

    for block in q_blocks:
        block = block.strip()
        if not block.startswith("**Question"):
            continue

        # Header line
        header_match = re.match(r'\*\*Question (\d+) --- Correct Answer: \(([a-d])\)\*\*', block)
        if not header_match:
            continue
        q_num, answer = header_match.groups()

        p = doc.add_paragraph(style="First Paragraph")
        run = p.add_run(f"Question {q_num} — Correct Answer: ({answer})")
        run.bold = True

        # Explanation and citation
        remaining = block[header_match.end():].strip()
        parts = remaining.split("\n\n")

        for part in parts:
            part = part.strip()
            if not part:
                continue
            if part.startswith("---"):
                continue

            # Citation line: *See: ...*
            if part.startswith("*See:") or part.startswith("*See :"):
                p = doc.add_paragraph(style="Body Text")
                run = p.add_run(part.strip("*"))
                run.italic = True
            else:
                # Explanation paragraph
                clean = part.replace("\n", " ")
                p = doc.add_paragraph(style="Body Text")
                # Handle inline italics
                segments = re.split(r'(\*[^*]+\*)', clean)
                for seg in segments:
                    if seg.startswith("*") and seg.endswith("*"):
                        run = p.add_run(seg[1:-1])
                        run.italic = True
                    else:
                        p.add_run(seg)

        # Separator
        doc.add_paragraph(style="Normal")

    doc.save(output_path)
    print(f"Student key saved to {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# Main — CLI interface
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate MCQ .docx deliverables from markdown drafts."
    )
    parser.add_argument("source_dir",
        help="Directory containing draft_full_set.md, draft_answer_key_full.md, "
             "and draft_answer_key_student.md")
    parser.add_argument("--output", default=DOWNLOADS,
        help="Output directory (default: ~/Downloads/)")
    parser.add_argument("--prefix", default="MCQ",
        help="Filename prefix (default: MCQ)")
    parser.add_argument("--school",
        default="UNIVERSITY OF PENNSYLVANIA CAREY LAW SCHOOL")
    parser.add_argument("--course",
        default="INTRODUCTION TO INTELLECTUAL PROPERTY \u2014 LAW 507")
    parser.add_argument("--semester", default="SPRING 2026")
    parser.add_argument("--professor", default="[Your Name]")
    parser.add_argument("--header",
        help="Custom header text for exam document")

    args = parser.parse_args()
    src = args.source_dir
    out = args.output
    pfx = args.prefix

    exam_md = os.path.join(src, "draft_full_set.md")
    ak_full_md = os.path.join(src, "draft_answer_key_full.md")
    ak_student_md = os.path.join(src, "draft_answer_key_student.md")

    for f in [exam_md, ak_full_md, ak_student_md]:
        if not os.path.exists(f):
            print(f"Error: {f} not found")
            exit(1)

    os.makedirs(out, exist_ok=True)
    meta = dict(school=args.school, course=args.course,
                semester=args.semester, professor=args.professor)

    print("Generating MCQ deliverables...\n")

    generate_exam(exam_md, os.path.join(out, f"{pfx}_Exam.docx"),
                  **meta, header_text=args.header)
    print()
    generate_answer_key(ak_full_md, os.path.join(out, f"{pfx}_AnswerKey.docx"),
                        **meta)
    print()
    generate_student_key(ak_student_md,
                         os.path.join(out, f"{pfx}_AnswerKey_Student.docx"),
                         **meta)

    print(f"\nDone. Files saved to {out}/")
